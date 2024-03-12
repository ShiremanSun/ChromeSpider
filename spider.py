#encoding: utf-8
from selenium import webdriver
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from datetime import datetime
import re
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Pt, RGBColor
import docx

def is_valid_date(date_text):
    try:
        datetime.strptime(date_text, '%Y-%m-%d')
        return True
    except ValueError:
        raise Exception("日期格式有误")

start_date = '2024-03-01'
end_date = '2024-03-12'

print("请输入开始时间，例如2024-03-01，然后按回车键")
start_date = input()
is_valid_date(start_date)
print("请输入结束时间，例如2024-03-06，然后按回车键")
end_date = input()
is_valid_date(end_date)
    
 



formate1 = "%Y-%m-%d"
formate2 = "%Y/%m/%d"


document = Document()
document.add_heading('政策', 1)

options = webdriver.ChromeOptions() 
options.add_argument('--headless')
browser = webdriver.Chrome(options=options)
WAIT = WebDriverWait(browser, 10)
browser.set_window_size(1400, 900)



class Item:
    def __init__(self, url, title, date):
        self.url = url
        self.title = title
        self.date = date
    def __str__(self):
        return ("链接: {}; 标题: {}; 日期:{}".format(self.url, self.title, self.date))
        
def vailidUrl(url):
    return url.startswith('http')
def checkDate(date):
    date = re.sub(r'[^0-9-/]', '', date)
    min = datetime.strptime(start_date, formate1)
    max = datetime.strptime(end_date, formate1)
    if '/' in date:
        format = formate2
    else:
        format = formate1
    return  min < datetime.strptime(date, format) and max > datetime.strptime(date, format)
def checkYear(year):
    return year <= 2024 and year >= 2023

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)
    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

order = 0
def appendDocument(province, data):
    global order
    order+=1
    #一级标题
    heading = document.add_heading(str(order) + '.' + province, 1)
    if len(data) == 0:
        document.add_paragraph('暂无数据', style='Heading 2')
        return
    for item in data:
        graph = document.add_paragraph(style='Heading 2')
        run = graph.add_run(item.title)
        run.font.color.rgb = RGBColor(0, 0, 0)
        add_hyperlink(graph, item.url, item.url)
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
def hebei():
    browser.get("https://kjt.hebei.gov.cn/www/xxgk2020/228104/228108/228109/index.html")
    host = 'https://kjt.hebei.gov.cn/'
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    uls = soup.find_all('ul', class_='xxgk-xxbox').pop(0)
    items = uls.find_all('li')
    results = list()
    for item in items: 
            date = item.find('span', class_='xxgk-line-comdate fr')
            herf = item.find('a').get('href')
            if not vailidUrl(herf):
                herf = urljoin(host, herf)
            if not checkDate(date.text):
                continue
            data = Item(herf, item.find('a').text, date.text)
            list.append(results, data)
            print(data)
    appendDocument("河北省",results)
    
def shanxi():
    results = list()
    browser.get("http://kjt.shanxi.gov.cn/zcfg/sxskjcxzc/")
    host = 'https://kjt.shanxi.gov.cn/'
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    div = soup.find('div', class_='rightcontent')
    items = div.find_all('li')

    for item in items: 
        a = item.find('a')
        link = a.get('href')
        last_ = link.rfind('_')
        lastT = link.rfind('t', 0, last_)
        date = link[lastT+1:last_]
        year = date[0:4]
        month = date[4:6]
        day = date[6:8]
        date_str = year + '-' + month + '-' + day
        if not checkDate(date_str):
                continue
        if not vailidUrl(link):
            link = urljoin(host, link)
        item = Item(link, a.text.replace("\n", "").replace('\t', ''), date_str)
        print(item)
        results.append(item)
    appendDocument("山西省",results)
def jilin():
    results = list()
    host = "http://kjt.jl.gov.cn/xxgk/fgwj/"
    browser.get(host)
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    table = soup.find_all('table', class_='paddingtop')[1] # 只取第一部分的
    
    try:
        tbody = table.find('tbody')
        trs = tbody.find_all('tr')
        if len(trs) <= 0:
            return
        trs.pop(0) #删除第一个
        for tr in trs:
            tds = tr.find_all('td')
            a = tds[0].find('a')
            if len(tds) >= 2:
                date = tds[1].text.replace('[','').replace(']',"")
                link = a.get('href')
                if not vailidUrl(link):
                    link = urljoin(host, link)
                if not checkDate(date):
                    continue
                item = Item(link, a.get('title'), date)
                print(item)
                results.append(item)
    except Exception as err:
        print(err)
    appendDocument("吉林省",results)      
        
def liaoning():
    results = list()
    host = "https://kjt.ln.gov.cn/kjt/kjzc/lnkjzc/index.shtml"
    browser.get(host)
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    ul = soup.find('ul', class_='govCListBox')
    try: 
        lis = ul.find_all('li')
        for li in lis: 
            date = li.find('span').text
            if not checkDate(date):
                continue
            link = li.find('a').get('href')
            if not vailidUrl(link):
                link = urljoin(host, link)
            item = Item(link, li.find('a').get('title'), date)
            print(item)
            results.append(item)
    except Exception as err:
        err.with_traceback()
        print(err)
    appendDocument("辽宁省",results)
    
def heilongjiang():
    results = list()
    host = "http://kjt.hlj.gov.cn/"
    browser.get("http://kjt.hlj.gov.cn/kjt/c113914/common_zfxxgk.shtml?tab=gkzc")
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    zc = soup.find('div', class_='zfxx-content zfxx-content-zc')
    tdbody = zc.find('tbody', id='zc-list-content2')
    trs = tdbody.find_all('tr')
    for tr in trs:
        title = tr.find('td', class_='info')
        date = tr.find_all('td').pop(4).text
        if not checkDate(date):
            continue
        link = title.find('a').get('href')
        if not vailidUrl(link):
            link = urljoin(host, link)
        item = Item(link, title.text, date)
        print(item)
        results.append(item)
    appendDocument("黑龙江省",results)

def gansu():
    results = list()
    browser.get("https://kjt.gansu.gov.cn/kjt/c111483/xxgk_infolist.shtml")
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    items = soup.find_all('ul', class_='kjtxxgklist')
    for item in items: 
        li = item.find('li')
        date = li.find('b')
        if not checkDate(date.text):
            continue
        item = Item(li.find('a').get('href'), li.find('a').get('title'), date.text)
        print(item)
        results.append(item)
    if(len(results) == 0):
        print("甘肃获取不到数据， 请手动获取")
    appendDocument("甘肃省",results)
def qinghai():
    results = list()
    browser.get("https://kjt.qinghai.gov.cn/content/lists/cid/22/pid/20/page/")
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    ul = soup.find('ul', class_='list_ul')
    items = ul.find_all('li')
    for item in items: 
        a = item.find('a')
        date = a.find('span')
        if not checkDate(date.text):
            continue
        item = Item(a.get('href'), a.text, date.text)
        print(item)
        results.append(item)
    appendDocument("青海省",results)
    
def shandong():
    results = list()
    host = "http://kjt.shandong.gov.cn"
    browser.get("http://kjt.shandong.gov.cn/col/col13349/index.html")
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    
    ul = soup.find('ul', id='policyRelease')
    items = ul.find_all('li')
    for item in items: 
        a = item.find('a')
        date = item.find('span')
        if not checkDate(date.text):
            continue
        link = a.get('href')
        if not vailidUrl(link):
            link = urljoin(host, link)
        item = Item(link, a.text.replace('...', ''), date.text)
        print(item)
        results.append(item)
    appendDocument("山东省",results)
def shanghai():
    results = list()
    host = "https://stcsm.sh.gov.cn/"
    browser.get("https://stcsm.sh.gov.cn/zwgk/kjzc/zcwj/#rightSidebar")
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    
    ul = soup.find('ul', class_='yjsnews')
    items = ul.find_all('li')
    for item in items: 
        a = item.find('a')
        date = item.find('span')
        if not checkDate(date.text):
                continue
        link = a.get('href')
        if not vailidUrl(link):
            link = urljoin(host, link)
        item = Item(link, a.get("title"), date.text)
        print(item)
        results.append(item)
    appendDocument("上海市",results)
    
def fujian():
    results = list()
    host = 'https://kjt.fujian.gov.cn/xxgk/zcwj/'
    browser.get(host)
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    
    ul = soup.find('div', class_='gl')
    items = ul.find_all('li')
    for item in items: 
        a = item.find('a')
        date = item.find('span').text
        link = a.get('href')
        if not checkDate(date):
                continue
        if not vailidUrl(link):
            link = urljoin(host, link)
        item = Item(link, a.get("title"), date)
        print(item)
        results.append(item)
    appendDocument("福建省",results)
def zhejiang():
    results = list()
    host = 'https://kjt.zj.gov.cn/'
    browser.get("https://kjt.zj.gov.cn/col/col1229080140/index.html")
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    
    ul = soup.find('div', class_='default_pgContainer')
    items = ul.find_all('li')
    for item in items: 
        a = item.find('a')
        date = item.find('span')
       
        if not checkDate(date.text):
            continue
        link = a.get('href')
        if not vailidUrl(link):
            link = urljoin(host, link)
        item = Item(link, a.get('title'), date.text)
        print(item)
        results.append(item)
    appendDocument("浙江省",results)
#todo 反爬
def henan():
    results = list()
    browser.get("https://kjt.henan.gov.cn/zcwj/")
    browser.implicitly_wait(10)
    html = browser.page_source
    try:
        soup = BeautifulSoup(html, 'html.parser')
        div = soup.find('div', class_='list_u1')
        items = div.find_all('li')
        for item in items: 
            a = item.find('a')
            date = item.find('span')
            if not checkDate(date.text):
                continue
            item = Item(a.get('href'), a.get('title'), date.text)
            print(item)
            results.append(item)
    except Exception as err:
        print(err)
        print("河南省数据获取失败，请手动获取")
    appendDocument("河南省",results)
    
def hubei():
    results = list()
    browser.get("https://kjt.hubei.gov.cn/zfxxgk_GK2020/zc2020/gfxwjj/#test")
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    ul = soup.find('ul', id='zcwj')
    items = ul.find_all('li')
    for item in items: 
        a = item.find('a')
        date = a.find('span')
        if not checkDate(date.text):
                continue
        item = Item(a.get('href'), a.get('title'), date.text)
        print(item)
        results.append(item)
    appendDocument("湖北省",results)
    
def jiangsu():
    results = list()
    host = 'https://kxjst.jiangsu.gov.cn'
    browser.get("https://kxjst.jiangsu.gov.cn/col/col82571/index.html")
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    try:
        ul = soup.find('div', class_='default_pgContainer')
        items = ul.find_all('li')
        for item in items: 
            a = item.find('a')
            date = item.find('span')
            if not checkDate(date.text):
                continue
            link = a.get('href')
            if not vailidUrl(link):
                link = urljoin(host, link)
            item = Item(link, a.get('title'), date.text)
            print(item)
            results.append(item)
    except Exception as err:
        print("江苏省数据获取失败，请手动获取")
        print(err)
    appendDocument("江苏省",results)
    
def anhui():
    results = list()
    browser.get("http://kjt.ah.gov.cn/public/column/21671?type=6&action=xinzheng")
    browser.implicitly_wait(10)
    html = browser.page_source
    
    soup = BeautifulSoup(html, 'html.parser')
    table = soup.find('table', class_='r_xh_b clearfix')
    trs = table.find_all('tr')
    for tr in trs: 
        if tr.find('td'):
            div = tr.find('div', class_="title")
            date = tr.find('p').find_all('span').pop(1).text
            a = div.find('a')
            if not checkDate(date):
                continue
            item = Item(a.get('href'), a.get('title'), date)
            print(item)
            results.append(item)
    appendDocument("安徽省",results)

def guangdong():
    results = list()
    browser.get("https://gdstc.gd.gov.cn/zwgk_n/zcfg/gfwj/index.html")
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    ul = soup.find('ul', class_='list')
    items = ul.find_all('li')
    for item in items: 
        a = item.find('a')
        date = item.find('span', class_='time')
        if not checkDate(date.text):
                continue
        item = Item(a.get('href'), a.text, date.text)
        print(item)
        results.append(item)
    appendDocument("广东省",results)
    
def hainan():
    results = list()
    host = "https://dost.hainan.gov.cn/xxgk/"
    browser.get("https://dost.hainan.gov.cn/xxgk/zfwj/")
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    ul = soup.find('ul', class_='zfwj-lis')
    items = ul.find_all('li')
    for item in items: 
        a = item.find('a')
        date = item.find('span', class_='time')
        if not checkDate(date.text):
                continue
        link = a.get('href')
        if not vailidUrl(link):
            link = link.replace('../', '')
            link = urljoin(host, link)
        item = Item(link, a.get('title'), date.text)
        print(item)
        results.append(item)
    appendDocument("海南省",results)
    
def sichuan():
    results = list()
    host = 'https://kjt.sc.gov.cn/'
    try:
        
        browser.get("https://kjt.sc.gov.cn/kjt/sjkjfg/newszwxxgkchild.shtml")
        browser.implicitly_wait(15)
        html = browser.page_source
        soup = BeautifulSoup(html, 'html.parser')
        div = soup.find('div', class_='content')
        ul = div.find('ul', class_='list-unstyled')
        items = ul.find_all('li')
        for item in items: 
            a = item.find('a')
            date = item.find('span')
            if not checkDate(date.text):
                continue
            link = a.get('href')
            if not vailidUrl(link):
                link = urljoin(host, link)
            item = Item(link, a.text, date.text)
            print(item)
            results.append(item)
    except Exception as err:
        print(err)
        print("四川获取失败，请手动获取")
    appendDocument("四川省",results)
    
def guizhou():
    
    results = list()
    browser.get("https://kjt.guizhou.gov.cn/zwgk/xxgkml/zcwj/gzhgfxwjsjk/gfxwjsjk/")
    browser.implicitly_wait(20)
    html = browser.page_source
    
    soup = BeautifulSoup(html, 'html.parser')
    try:
        table = soup.find('tbody', id='Result')
        trs = table.find_all('tr')
        for tr in trs: 
            tds = tr.find_all('td')
            a = tds.index(0).find('a')
            date = tds.index(2)
            if not checkDate(date.text):
                continue
            item = Item(a.get('href'), a.get('title'), date.text)
            print(item)
            results.append(item)
    except Exception as err:
        print(err)
        print("贵州数据获取失败，请手动获取")
    appendDocument("贵州省",results)
def yunnan():
    results = list()
    host = 'https://kjt.yn.gov.cn/'
    browser.get("https://kjt.yn.gov.cn/html/zhengfuxinxigongkai/zhengcewenjian/xingzhengguifanxingwenjian/")
    browser.implicitly_wait(10)
    html = browser.page_source
    
    soup = BeautifulSoup(html, 'html.parser')
    trs = soup.find_all('tr', class_="tml")
    for tr in trs: 
            td = tr.find('td', class_='tcc')
            a = td.find('a')
            date = td.find('i').text
            pattern = re.compile(r'〔(\d{4})〕')
            group = pattern.search(date)
            if group is None:
                continue
            year = int(group.group(1))
            
            if not checkYear(year):
                continue
            print(year)
            link = a.get('href')
            if not vailidUrl(link):
                link = urljoin(host, link)
            item = Item(link, a.text, year)
            print(item)
            results.append(item)
    print('云南无法获取月份，请double check')
    appendDocument("云南省",results)
def jiangxi():
    results = list()
    browser.get("https://kjt.jiangxi.gov.cn/col/col64258/index.html#div")
    browser.implicitly_wait(20)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    try:
        table = soup.find('tbody')
        trs = table.find_all('tr')
        trs.pop(0)
        for tr in trs: 
            tds = tr.find_all('td')
            a = tds[0].find('a')
            date = tds[3]
            if not checkDate(date.text):
                    continue
            item = Item(a.get('href'), a.get('title'), date.text)
            print(item)
            results.append(item)
    except Exception as err:
        print("江西获取失败，请重试")
    appendDocument("江西省",results)
    
def hunan():
    results = list()
    host = 'https://kjt.hunan.gov.cn/'
    browser.get("https://kjt.hunan.gov.cn/kjt/xxgk/zcfg/tgfxwj/index.html")
    browser.implicitly_wait(10)
    html = browser.page_source
    
    soup = BeautifulSoup(html, 'html.parser')
    div = soup.find('div', class_='xzadd-mu-1')
    table = div.find('tbody')
    trs = table.find_all('tr')
    for tr in trs: 
            tds = tr.find_all('td')
            a = tds[1].find('a')
            date = tds[3]
            if not checkDate(date.text):
                continue
            link = a.get('href')
            if not vailidUrl(link):
                link = urljoin(host, link)
            item = Item(link, a.text, date.text)
            print(item)
            results.append(item)
    appendDocument("湖南省",results)
def tianjin():
    results = list()
    host = 'https://kxjs.tj.gov.cn/'
    browser.get("https://kxjs.tj.gov.cn/ZWGK4143/ZCFG148_1/ZCFB4222/2023Nzcfb_160567/")
    browser.implicitly_wait(10)
    html = browser.page_source
    
    soup = BeautifulSoup(html, 'html.parser')
    ul = soup.find('ul', class_='news_list news_list2')
    items = ul.find_all('li')
    for item in items: 
        a = item.find('a')
        date = a.find('span', class_='time')
        if not checkDate(date.text):
                continue
        link = a.get('href')
        if not vailidUrl(link):
            link = urljoin(host, link)
        item = Item(link, a.text, date.text)
        print(item)
        results.append(item)
    appendDocument("天津市",results)
    
def beijing():
    results = list()
    browser.get("https://kw.beijing.gov.cn/col/col2384/index.html")
    browser.implicitly_wait(10)
    html = browser.page_source
    host = 'https://kw.beijing.gov.cn'
    soup = BeautifulSoup(html, 'html.parser')
    td = soup.find('td', id='newslist_4812_4812_4812')
    try:
        trs = td.find_all('tr')
        for tr in trs: 
            tds = tr.find_all('td')
            a = tds[0].find('a')
            date = tds[1].find('font')
            if not checkDate(date.text):
                continue
            link = a.get('href')
            if not vailidUrl(link):
                link = urljoin(host, link)
            item = Item(link, a.text, date.text)
            print(item)
            results.append(item)
    except Exception as err:
        print("北京获取错误,请手动获取")
        print(err)
    appendDocument("北京市",results)
    
def neimeng():
    results = list()
    browser.get("https://kjt.nmg.gov.cn/zwgk/zfxxgk/fdzdgknr/?gk=2&cid=11919#iframe")
    browser.implicitly_wait(10)
    html = browser.page_source
    
    soup = BeautifulSoup(html, 'html.parser')
    table = soup.find_all('table', id='table1')[1]
    trs = table.find('tbody').find_all('tr')
    for tr in trs: 
            tds = tr.find_all('td')
            date = tds[5]
            a = tds[1].find('a')
            if not checkDate(date.text):
                continue
            item = Item(a.get('href'), a.text, date.text)
            print(item)
            results.append(item)
    appendDocument("内蒙古",results)
def guangxi():
    results = list()
    host = 'http://kjt.gxzf.gov.cn/xxgk/zcyd/zcfg/'
    browser.get(host)
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    ul = soup.find('ul', class_='more-list')
    items = ul.find_all('li')
    for item in items: 
        a = item.find('a')
        date = item.find('span')
        if not checkDate(date.text):
                continue
        link = a.get('href')
        if not vailidUrl(link):
            link = urljoin(host, link)
        item = Item(link, a.get('title'), date.text)
        print(item)
        results.append(item)
    appendDocument("广西",results)
def xizang():
    results = list()
    host = 'https://sti.xizang.gov.cn/xxgk/gfxwj/'
    browser.get(host)
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    ul = soup.find('ul', class_='gl-l')
    items = ul.find_all('li')
    for item in items: 
        a = item.find('a')
        date = item.find('span')
        if not checkDate(date.text):
                continue
        link = a.get('href')
        if not vailidUrl(link):
            link = urljoin(host, link)
        item = Item(link, a.text, date.text)
        print(item)
        results.append(item)
    appendDocument("西藏",results)
    
def ningxia():
    results = list()
    host = 'https://kjt.nx.gov.cn/zcfg/tfwj/'
    browser.get(host)
    browser.implicitly_wait(10)
    html = browser.page_source
    
    soup = BeautifulSoup(html, 'html.parser')
    ul = soup.find('ul', class_='list_ul fs')
    items = ul.find_all('li')
    for item in items: 
        a = item.find('a')
        date = item.find('span')
        if not checkDate(date.text):
                continue
        link = a.get('href')
        if not vailidUrl(link):
            link = urljoin(host, link)
        item = Item(link, a.text.replace('》', ''), date.text)
        print(item)
        results.append(item)
    appendDocument("宁夏",results)
def chongqing():
    results = list()
    host = 'https://kjj.cq.gov.cn/zwgk_176/zwxxgkml/zcwj/xzgfxwj/'
    browser.get(host)
    browser.implicitly_wait(10)
    html = browser.page_source
    
    soup = BeautifulSoup(html, 'html.parser')
    table = soup.find('table', class_='zcwjk-list')
    trs = table.find('tbody').find_all('tr')
    trs.pop(0)
    for tr in trs:
            data = tr.find('td', class_='title')
            a = data.find('a')
            link = a.get('href')
            last_ = link.rfind('_')
            lastT = link.rfind('t', 0, last_)
            date = link[lastT+1:last_]
            year = date[0:4]
            month = date[4:6]
            day = date[6:8]
            date_str = year + '-' + month + '-' + day
            if not checkDate(date_str):
                continue
            link = a.get('href')
            if not vailidUrl(link):
                link = urljoin(host, link)
            item = Item(link, a.find('p').text, date_str)
            print(item)
            results.append(item)
    appendDocument("重庆",results)
def shaanxi():
    results = list()
    host = 'https://kjt.shaanxi.gov.cn'
    browser.get("https://kjt.shaanxi.gov.cn/zcwj.html")
    browser.implicitly_wait(10)
    html = browser.page_source
    
    soup = BeautifulSoup(html, 'html.parser')
    div = soup.findAll('div', class_='gbx')[3]

    ul = div.find('ul')
    items = ul.find_all('li')
    for item in items: 
        a = item.find('a')
        date = item.find('span')
        if not checkDate(date.text):
                continue
        link = a.get('href')
        if not vailidUrl(link):
            link = urljoin(host, link)
        item = Item(link, a.get('title'), date.text)
        print(item)
        results.append(item)
    appendDocument("陕西",results)
def xinjiang():
    results = list()
    host = 'https://kjt.xinjiang.gov.cn/'
    browser.get("https://kjt.xinjiang.gov.cn/kjt/c100271/zfxxgk_gknrz.shtml")
    browser.implicitly_wait(10)
    html = browser.page_source
    soup = BeautifulSoup(html, 'html.parser')
    div = soup.find('div', class_='gknr_list')
    items = div.find_all('dd')
    for item in items: 
        a = item.find('a')
        date = item.find('span')
        if not checkDate(date.text):
                continue
        link = a.get('href')
        if not vailidUrl(link):
            link = urljoin(host, link)
       
        item = Item(link, a.text, date.text)
        print(item)
        results.append(item)
    appendDocument("新疆",results)
def shenzhen():
    results = list()
    browser.get("http://stic.sz.gov.cn/xxgk/zcfg/zcfg/")
    browser.implicitly_wait(10)
    html = browser.page_source
    
    soup = BeautifulSoup(html, 'html.parser')
    ul = soup.find('ul', id='fgList')
    items = ul.find_all('li')
    for item in items: 
        a = item.find('a')
        date = item.find('span')
        if not checkDate(date.text):
            continue
        item = Item(a.get('href'), a.text, date.text)
        print(item)
        results.append(item)
    appendDocument("深圳",results)
hebei()
shanxi()
jilin()
liaoning()
heilongjiang()
gansu() # 获取不到
qinghai()
shandong()
shanghai()
fujian()
zhejiang()
henan() #反爬
hubei()
jiangsu()
anhui()
guangdong()
hainan()
sichuan()
guizhou() #拿不到数据
yunnan() #拿不到月份
jiangxi() #有概率失败
hunan()
tianjin()
beijing()
neimeng()
guangxi()
xizang()
ningxia()
chongqing()
shaanxi()
xinjiang()
shenzhen()

document.save('test.docx')

